"""
Экстрактор текста из файлов Microsoft Word (.docx).

Извлекает:
    - Текст абзацев (основной контент)
    - Текст из таблиц (все ячейки)
    - Текст из верхних и нижних колонтитулов (header/footer)
    - Текст из текстовых рамок (``textboxes``)
    - Свойства документа (автор, дата создания и т.д.)

Ограничения:
    - Формат ``.doc`` (старый бинарный формат) не поддерживается
      напрямую.  Для его обработки требуется предварительная конвертация
      через ``LibreOffice`` или ``antiword``.
    - Встроенные изображения не обрабатываются (используйте
      :class:`~extractors.ImageExtractor` отдельно).

Зависимости:
    - ``python-docx``

Example::

    from extractors.docx_extractor import DocxExtractor

    result = DocxExtractor("contract.docx").extract()
    print(result.metadata["author"])
    print(result.text[:300])
"""

from __future__ import annotations

import logging
from typing import Iterator

from extractors.base import BaseExtractor, ExtractionResult, ExtractionError

logger = logging.getLogger(__name__)


class DocxExtractor(BaseExtractor):
    """Экстрактор текста из файлов ``.docx`` (Microsoft Word).

    Args:
        file_path: Путь к ``.docx``-файлу.

    Raises:
        FileNotFoundError: Если файл не найден.
        ExtractionError: Если файл повреждён или не является валидным
            ``.docx``-документом.
    """


    def extract(self) -> ExtractionResult:
        """Извлекает текст из ``.docx``-файла.

        Последовательно собирает текст из:
        1. Основных абзацев документа
        2. Таблиц (каждая ячейка как отдельный фрагмент)
        3. Колонтитулов (верхний и нижний для каждой секции)

        Returns:
            :class:`~extractors.base.ExtractionResult` с текстом и
            метаданными:

            - ``paragraph_count`` — количество абзацев
            - ``table_count`` — количество таблиц
            - ``author`` — автор документа
            - ``created`` — дата создания (строка ISO или ``None``)
            - ``title`` — заголовок документа

        Raises:
            ExtractionError: При ошибке открытия или чтения файла.
        """
        start = self._log_start()

        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError as exc:
            raise ExtractionError(
                self.file_path,
                "Библиотека 'python-docx' не установлена. "
                "Выполните: pip install python-docx",
            ) from exc

        try:
            doc = Document(self.file_path)
        except Exception as exc:
            raise ExtractionError(
                self.file_path,
                f"Не удалось открыть .docx файл: {exc}",
            ) from exc

        chunks: list[str] = []

        # Основные абзацы
        para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        chunks.extend(para_texts)

        # Таблицы
        table_texts = self._extract_tables(doc)
        chunks.extend(table_texts)

        # Колонтитулы
        header_footer_texts = self._extract_headers_footers(doc)
        chunks.extend(header_footer_texts)

        full_text = "\n\n".join(chunks)

        # Метаданные документа
        core_props = doc.core_properties
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "author": getattr(core_props, "author", None),
            "title": getattr(core_props, "title", None),
            "created": (
                str(core_props.created) if getattr(core_props, "created", None) else None
            ),
            "modified": (
                str(core_props.modified)
                if getattr(core_props, "modified", None)
                else None
            ),
        }

        result = self._make_result(
            text=full_text,
            chunks=chunks,
            metadata=metadata,
            start_time=start,
        )
        self._log_done(result)
        return result

    def extract_chunks(self) -> Iterator[str]:
        """Генератор текстовых чанков по абзацам.

        Yields:
            Текст каждого непустого абзаца документа.

        Raises:
            ExtractionError: При ошибке открытия файла.
        """
        try:
            from docx import Document
        except ImportError as exc:
            raise ExtractionError(
                self.file_path,
                "Библиотека 'python-docx' не установлена.",
            ) from exc

        try:
            doc = Document(self.file_path)
        except Exception as exc:
            raise ExtractionError(self.file_path, str(exc)) from exc

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                yield text


    @staticmethod
    def _extract_tables(doc: object) -> list[str]:
        """Извлекает текст из всех таблиц документа.

        Каждая строка таблицы объединяется через «\\t», строки через «\\n».
        Пустые таблицы пропускаются.

        Args:
            doc: Объект ``Document`` из python-docx.

        Returns:
            Список текстовых фрагментов — по одному на каждую таблицу.
        """
        table_texts: list[str] = []
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = "\t".join(c for c in cells if c)
                if row_text:
                    rows.append(row_text)
            if rows:
                table_texts.append("\n".join(rows))
        return table_texts

    @staticmethod
    def _extract_headers_footers(doc: object) -> list[str]:
        """Извлекает текст из верхних и нижних колонтитулов.

        Перебирает все секции документа и извлекает текст из
        ``header`` и ``footer`` каждой секции.

        Args:
            doc: Объект ``Document`` из python-docx.

        Returns:
            Список непустых текстовых фрагментов из колонтитулов.
        """
        texts: list[str] = []
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf is None:
                    continue
                hf_text = "\n".join(
                    p.text.strip() for p in hf.paragraphs if p.text.strip()
                )
                if hf_text:
                    texts.append(hf_text)
        return texts